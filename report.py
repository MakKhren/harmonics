from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Pt


class Report():
    def __init__(self, my_tnd_i, my_tnd_u, name_files):
        self.my_tnd_i = my_tnd_i
        self.my_tnd_u = my_tnd_u
        self.name_files = name_files
        try:
            self.doc = Document(self.name_files)
        except:
            self.doc = Document()
        self.__heading_text = [
            'Исходные данные',
            'Анализ гармонических искажений',
        ]
        self.__para_text =[
            f'- полный расчетный действующий ток на присоединении '
            f'Iд = {self.my_tnd_i.I_oper_full} А ',
            f'- расчетный действующий ток преобразовательной техники '
            f'Iпт = {self.my_tnd_i.I_oper_conv} А',
            f'- количество гармоник учитывающийся при расчете '
            f'n = {self.my_tnd_i.n_harm} ',
            f'- действующие значение фазного напряжения '
            f'Uф = {self.my_tnd_u.U_phase_oper} В',
            f'- активно сопротивление системы '
            f'Rпс = {self.my_tnd_u.r_sys} ',
            f'- реактивное сопротивление системы '
            f'Xпс = {self.my_tnd_u.x_sys} ',
            'Таблица 1 Амплитудные значения токов гармоник', 
            'Среднеквадратичное значение тока гармоник:',
            f'Irms = {my_tnd_i.get_I_rms_harmonics()} A',
            'Коэффициент искажения синусоидальности тока: ',
            f'TNDi = {my_tnd_i.get_tnd_i()} %', 
            'Таблица 2 Амплитудные значения напряжений гармоник',
            'Среднеквадратичное значение напряжения гармоник:',
            f'Urms ={my_tnd_u.get_U_rms_harmonics()} В',
            'Коэффициент искажения синусоидальности напряжения:',
            f'TNDu = {my_tnd_u.get_tnd_u()} %',
        ]
        self.__table_cols_text = {
            'table_1': ['n', 'InA, А'],
            'table_2': ['n', 'bn1, B', 'bn2, B', 'bn3, B', 'Un, В'],
        }
        self.__value_I = [
            range(1,len(self.my_tnd_i.get_I_amp_harmonics())+1),
            self.my_tnd_i.get_I_amp_harmonics(),
            ]
        self.__value_U = [
            range(1,len(self.my_tnd_u.get_U_amp_harmonics())+1),
            self.my_tnd_u.get_U_amp_harmonics_bn1(),
            self.my_tnd_u.get_U_amp_harmonics_bn2(),
            self.my_tnd_u.get_U_amp_harmonics_bn3(),
            self.my_tnd_u.get_U_amp_harmonics(),
        ]

    def __get_para_obj (self, text, aling='left', space=False):
        para_obj = self.doc.add_paragraph(text)
        para_obj.paragraph_format.first_line_indent = Pt(24)
        para_obj.runs[0].font.name = 'Arial'
        para_obj.runs[0].font.size = Pt(12)
        para_obj.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        if aling == 'center':
            para_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_obj.paragraph_format.space_before =  Pt(6)
            para_obj.paragraph_format.space_after =  Pt(6)
        elif aling == 'left':
            para_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if space == True:
                para_obj.paragraph_format.space_before =  Pt(6)
                para_obj.paragraph_format.space_after =  Pt(0)
            else:
                para_obj.paragraph_format.space_before =  Pt(0)
                para_obj.paragraph_format.space_after =  Pt(0)

    def __get_heading_obj(self, text,aling='left'):
        heading_obj =self.doc.add_heading(text, 2)
        heading_obj.runs[0].font.name = 'Arial'
        heading_obj.runs[0].font.size = Pt(14)
        heading_obj.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        heading_obj.runs[0].font.bold = False
        heading_obj.paragraph_format.space_after =  Pt(6)
        heading_obj.paragraph_format.space_before =  Pt(6)
        if aling == 'center':
            heading_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif aling == 'left':
           heading_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def __get_headlines_table(self, name_cols, table):
        for num_col, name_col in enumerate(name_cols):
            table.rows[0].cells[num_col].text = name_col
            cell_para = table.rows[0].cells[num_col].paragraphs[0]
            cell_para.runs[0].font.name = 'Arial'
            cell_para.runs[0].font.size = Pt(12)
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __get_all_row(self, list_values, table):
        for row in range(len(list_values)):
            row_cells = table.add_row().cells

    def __write_values_table(self, list_values, table):
        for col, values in enumerate(list_values):
            for row, value  in enumerate(values):
                table.rows[row+1].cells[col].text = str(value)
                cell_para = table.rows[row+1].cells[col].paragraphs[0]
                cell_para.runs[0].font.name = 'Arial'
                cell_para.runs[0].font.size = Pt(12)
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __get_table(self, name_cols, list_values):
        table = self.doc.add_table(rows=1, cols=len(name_cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        self.__get_headlines_table(name_cols, table)
        self.__get_all_row(list_values[0], table)
        self.__write_values_table(list_values, table)

    def get_report(self):
        self.__get_heading_obj(self.__heading_text[0])
        self.__get_para_obj(self.__para_text[0])
        self.__get_para_obj(self.__para_text[1])
        self.__get_para_obj(self.__para_text[2])
        self.__get_para_obj(self.__para_text[3])
        self.__get_para_obj(self.__para_text[4])
        self.__get_para_obj(self.__para_text[5])
        self.__get_heading_obj(self.__heading_text[1])
        self.__get_para_obj(self.__para_text[6])
        self.__get_table(self.__table_cols_text['table_1'], self.__value_I)
        self.__get_para_obj(self.__para_text[7], space=True)
        self.__get_para_obj(self.__para_text[8], 'center')
        self.__get_para_obj(self.__para_text[9])
        self.__get_para_obj(self.__para_text[10], 'center')
        self.__get_para_obj(self.__para_text[11])
        self.__get_table(self.__table_cols_text['table_2'], self.__value_U)
        self.__get_para_obj(self.__para_text[12], space=True)
        self.__get_para_obj(self.__para_text[13], 'center')
        self.__get_para_obj(self.__para_text[14])
        self.__get_para_obj(self.__para_text[15], 'center')
        if self.my_tnd_i.get_tnd_i() < 10 and self.my_tnd_u.get_tnd_u() < 8:
            text = 'Применение активных фильтров гармоник не требуется.'
            self.__get_para_obj(text)
        else:
            text = ('Необходимо применение активных фильтров гармоник для '
            'компенсации токов гармоник в соответсвии с Таблицей 1')
            self.__get_para_obj(text)
        self.doc.save('report.docx')         
            
def print_consol(my_tnd_i, my_tnd_u):
    print()
    print(f'РЕЗУЛЬТАТЫ РАСЧЕТА'.center(50, '-'))
    print(f"TNDi = {my_tnd_i.get_tnd_i()} "
    f"TNDu = {my_tnd_u.get_tnd_u()}".center(50, '-'))
    print() 
    if my_tnd_i.get_tnd_i() < 10 and my_tnd_u.get_tnd_u() < 8:
            text = ('Применение активных фильтров гармоник '
            'не требуется.'.center(50, '-'))
            print(text)
    else:
        text = ('Необходимо применение активных фильтров гармоник для '
        'компенсации токов гармоник').center(50, '-')
        print(text)
    print() 